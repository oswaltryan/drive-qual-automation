from __future__ import annotations

import importlib
import json
from pathlib import Path
from typing import Any

import report_generation_helpers as h


def test_generate_report_docx_matches_reference_section_shape() -> None:
    from docx import Document

    source_root = Path("tests/.tmp/test_report_generation_shape")
    h._prepare_report_generation_shape_fixture(source_root)

    module = importlib.import_module("drive_qual.reports.generate")
    output_path = module.generate_report_docx(part_number="69-420", source_root=source_root)
    document = Document(output_path)
    headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style is not None and paragraph.style.name.startswith("Heading")
    ]
    table_headers = [" | ".join(cell.text for cell in table.rows[0].cells) for table in document.tables]
    expected_base_headings = [
        "Drive Qualification Report",
        "Executive Summary",
        "Drive Info",
        "Qualification Equipment",
        "Power Data",
        "Compatibility Data",
        "Disk Performance",
        "Compliance/Reliability Test",
        "Temperature Data",
    ]
    assert headings[: len(expected_base_headings)] == expected_base_headings
    assert h._appendix_headings(headings) == ["Disk Performance Raw Data & Measurements (Padlock DT)"]
    assert "Power Data" in headings
    assert "Compatibility Data" in headings
    assert "Disk Performance" in headings
    h._assert_reordered_sections(document, headings, "Padlock DT")
    h._assert_removed_headings_are_absent(headings)
    assert h._heading_starts_after_page_break(document, "Power Data")
    assert table_headers[:6] == [
        "Revision | Name | Date | Description",
        "Test | Linux | macOS | Windows",
        "Test | Linux | macOS | Windows",
        "DUT | CDM-R | CDM-W | BM(R) | BM(W) | ATTO-R | ATTO-W",
        "Program | Iterations/Loops | Result",
        "Chart | Temperature | Read MB/s | Write MB/s",
    ]
    assert h._table_contains_text(document.tables[2], "Native Disk Utility")
    assert not h._table_contains_text(document.tables[2], "Appears in Device Manager & Disk Management")
    assert h._table_contains_text(document.tables[4], "USB-IF Mass Storage Compliance")
    assert not h._table_contains_text(document.tables[4], "USB-IF Mass Storage Compliance (MSC)")
    summary = next(
        paragraph for paragraph in document.paragraphs if paragraph.text.startswith("Results require review")
    )
    assert summary.text == "Results require review in sections: Power Data."
    assert [run.text for run in summary.runs if run.bold] == ["Power Data"]
    h._assert_result_table_alignment(document)
    assert "Artifact | Path" not in table_headers
    assert h._has_paragraph_between_tables(document, "Windows | ", "Linux | ")
    assert h._has_paragraph_between_tables(document, "Linux | ", "macOS | ")
    assert h._first_column_is_narrower_than_artifact_column(document.tables[6])


def test_generate_report_docx_embeds_appendix_images_instead_of_paths() -> None:
    from docx import Document

    source_root = Path("tests/.tmp/test_report_generation_images")
    windows_dir = h._prepare_report_generation_images_fixture(source_root)

    module = importlib.import_module("drive_qual.reports.generate")
    output_path = module.generate_report_docx(part_number="69-420", source_root=source_root)
    document = Document(output_path)

    h._assert_appendix_object_layout(document, output_path)
    assert h._has_page_break_between_tables(document, "Linux | ", "macOS | ")
    assert "Artifact | Path" not in h._table_headers(document)
    assert str(windows_dir) not in h._document_text(document)
    assert "Padlock DT CrystalDiskInfo Drive Information.json" not in h._document_text(document)
    assert h._nested_table_rows(document.tables[6].rows[1].cells[1]) == [
        ["Name", "Measurement", "Accum-Mean", "Accum-Min", "Accum-Max"],
        ["Meas1", "Maximum", "448.48 mA", "444.94 mA", "453.56 mA"],
        ["Meas3", "RMS", "258.60 mA", "258.04 mA", "259.17 mA"],
        ["Meas8", "Max IO Excluded", "222.00 mA", "221.00 mA", "223.00 mA"],
    ]
    assert h._nested_table_rows(document.tables[6].rows[2].cells[1]) == [
        ["Name", "Measurement", "Accum-Mean", "Accum-Min", "Accum-Max"],
        ["Meas1", "Maximum", "448.48 mA", "444.94 mA", "453.56 mA"],
        ["Meas3", "RMS", "258.60 mA", "258.04 mA", "259.17 mA"],
        ["Meas6", "Inrush Excluded", "111.00 mA", "110.00 mA", "112.00 mA"],
    ]
    assert h._nested_table_columns_evenly_fill_parent(document.tables[6].rows[2].cells[1])
    assert h._table_columns_are_centered(document.tables[6].rows[2].cells[1].tables[0], range(2, 5))
    assert h._nested_tables_rows(document.tables[6].rows[3].cells[1]) == [
        ["ATTO", "", ""],
        ["I/O Size", "Write", "Read"],
        ["4KB", "120.00", "130.00"],
        ["64KB", "240.00", "250.00"],
        ["1MB", "320.00", "330.00"],
        ["4MB", "360.00", "370.00"],
        ["Crystal Disk Mark", ""],
        ["Metric", "Value"],
        ["Read MB/s", "350.93"],
        ["Write MB/s", "345.04"],
    ]
    assert h._table_columns_are_centered(document.tables[6].rows[3].cells[1].tables[0], h._value_column_indexes)
    assert h._table_columns_are_centered(document.tables[6].rows[3].cells[1].tables[1], h._value_column_indexes)
    assert document.tables[7].rows[2].cells[1].text == "Linux\\Padlock DT Max IO Summary.csv"
    h._assert_linux_disks_summary_table(document.tables[7])
    assert h._nested_tables_rows(document.tables[8].rows[3].cells[1]) == [
        ["Blackmagic", "", ""],
        ["Metric", "Read", "Write"],
        ["Speed", "350.93", "345.04"],
    ]
    assert h._table_columns_are_centered(document.tables[8].rows[3].cells[1].tables[0], h._value_column_indexes)
    assert h._nested_table_rows(document.tables[6].rows[4].cells[1]) == [
        ["Field", "Value"],
        ["Model", "Apricorn Padlock DT"],
        ["Transfer Mode", "SATA/600 | SATA/600"],
        ["Standard", "ACS-4"],
        ["Features", "S.M.A.R.T., NCQ, TRIM"],
        ["Rotation Rate", "---- (SSD)"],
    ]
    cdi_table = document.tables[6].rows[4].cells[1].tables[0]
    assert not h._table_column_paragraphs_are_centered(cdi_table, 0)
    assert h._table_column_paragraphs_are_centered(cdi_table, 1)


def test_generate_report_docx_formats_boolean_compliance_results_as_pass_fail() -> None:
    from docx import Document

    source_root = Path("tests/.tmp/test_report_generation_boolean_compliance")
    h._prepare_report_generation_shape_fixture(source_root)
    report_path = source_root / "69-420" / "drive_qualification_report_atomic_tests.json"
    data = json.loads(report_path.read_text(encoding="utf-8"))
    data["compliance"] = {
        "usb_if_msc_iterations": 4,
        "usb_if_msc_result": True,
        "disk_tester_reliability_iterations": 8,
        "disk_tester_reliability_result": False,
    }
    report_path.write_text(json.dumps(data), encoding="utf-8")

    module = importlib.import_module("drive_qual.reports.generate")
    output_path = module.generate_report_docx(part_number="69-420", source_root=source_root)
    document = Document(output_path)
    compliance_table = document.tables[4]

    assert [cell.text for cell in compliance_table.rows[1].cells] == [
        "USB-IF Mass Storage Compliance",
        "4",
        "Pass",
    ]
    assert [cell.text for cell in compliance_table.rows[2].cells] == ["Reliability Test", "8", "Fail"]
    assert _cell_shading_fill(compliance_table.rows[1].cells[2]) == "C6EFCE"
    assert _cell_shading_fill(compliance_table.rows[2].cells[2]) == "FFC7CE"


def test_generate_report_docx_marks_missing_cdi_details_red() -> None:
    from docx import Document

    source_root = Path("tests/.tmp/test_report_generation_missing_cdi")
    h._prepare_report_generation_shape_fixture(source_root)
    report_path = source_root / "69-420" / "drive_qualification_report_atomic_tests.json"
    data = json.loads(report_path.read_text(encoding="utf-8"))
    data["performance"]["Padlock DT"]["Windows"]["CrystalDiskInfo"] = {
        "screenshot": True,
        "model": "Apricorn Padlock DT",
    }
    report_path.write_text(json.dumps(data), encoding="utf-8")

    module = importlib.import_module("drive_qual.reports.generate")
    output_path = module.generate_report_docx(part_number="69-420", source_root=source_root)
    document = Document(output_path)
    summary = next(
        paragraph for paragraph in document.paragraphs if paragraph.text.startswith("Results require review")
    )
    cdi_table = document.tables[6].rows[4].cells[1].tables[0]

    assert "Disk Performance Raw Data & Measurements (Padlock DT)" in summary.text
    assert cdi_table.rows[1].cells[1].text == "Apricorn Padlock DT"
    assert _cell_shading_fill(cdi_table.rows[1].cells[1]) is None
    assert cdi_table.rows[2].cells[1].text == ""
    assert _cell_shading_fill(cdi_table.rows[2].cells[1]) == "FFC7CE"


def test_generate_report_docx_keeps_non_dt_linux_and_macos_appendix_tables_together() -> None:
    from docx import Document

    source_root = Path("tests/.tmp/test_report_generation_non_dt_appendix")
    h._prepare_report_generation_images_fixture(source_root, dut_name="Padlock SSD")

    module = importlib.import_module("drive_qual.reports.generate")
    output_path = module.generate_report_docx(part_number="69-420", source_root=source_root)
    document = Document(output_path)

    assert h._has_paragraph_between_tables(document, "Linux | ", "macOS | ")
    assert not h._has_page_break_between_tables(document, "Linux | ", "macOS | ")


def _cell_shading_fill(cell: Any) -> str | None:
    from docx.oxml.ns import qn

    shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    return None if shading is None else shading.get(qn("w:fill"))
