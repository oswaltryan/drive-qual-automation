from __future__ import annotations

import builtins
import importlib
import json
import sys
import zipfile
from pathlib import Path
from typing import Any, cast

from _pytest.monkeypatch import MonkeyPatch

from drive_qual.reports.evaluation import Status, case_material_for_product, evaluate_power, evaluate_temperature

EXPECTED_INLINE_IMAGE_COUNT = 1
EXPECTED_POWER_OBJECT_COUNT = 2
EXPECTED_MEDIA_FILE_COUNT = 3
MAX_WORD_OBJECT_ID = 2_000_000_000


def test_report_generation_import_does_not_import_docx(monkeypatch: MonkeyPatch) -> None:
    sys.modules.pop("drive_qual.reports.generate", None)

    real_import = builtins.__import__

    def guarded_import(
        name: str,
        globals: dict[str, Any] | None = None,
        locals: dict[str, Any] | None = None,
        fromlist: tuple[str, ...] = (),
        level: int = 0,
    ) -> Any:
        if name == "docx" or name.startswith("docx."):
            raise AssertionError("report generation imported python-docx eagerly")
        return real_import(name, globals, locals, fromlist, level)

    monkeypatch.setattr(builtins, "__import__", guarded_import)

    module = importlib.import_module("drive_qual.reports.generate")

    assert module.DEFAULT_OUTPUT_NAME == "drive_qualification_report.docx"


def test_generate_report_uses_source_root_for_input_and_default_output(monkeypatch: MonkeyPatch) -> None:
    module = importlib.import_module("drive_qual.reports.generate")
    source_root = Path("tests/.tmp/test_report_generation")
    part_dir = source_root / "69-420"
    part_dir.mkdir(parents=True, exist_ok=True)
    report_path = part_dir / "drive_qualification_report_atomic_tests.json"
    report_path.write_text(json.dumps({"drive_info": {"apricorn_part_number": "69-420"}}), encoding="utf-8")
    captured: dict[str, Any] = {}

    def fake_write_docx_report(
        data: dict[str, Any], evaluated: Any, source_report_path: Path, output_path: Path
    ) -> None:
        captured["data"] = data
        captured["source_report_path"] = source_report_path
        captured["output_path"] = output_path

    monkeypatch.setattr(module, "write_docx_report", fake_write_docx_report)

    output_path = module.generate_report_docx(part_number="69-420", source_root=source_root)

    assert output_path == part_dir / "drive_qualification_report.docx"
    assert captured["source_report_path"] == report_path
    assert captured["output_path"] == output_path
    assert captured["data"]["drive_info"]["apricorn_part_number"] == "69-420"


def test_generate_report_docx_matches_reference_section_shape() -> None:
    from docx import Document

    source_root = Path("tests/.tmp/test_report_generation_shape")
    part_dir = source_root / "69-420"
    part_dir.mkdir(parents=True, exist_ok=True)
    report_path = part_dir / "drive_qualification_report_atomic_tests.json"
    report_path.write_text(json.dumps(_report_payload()), encoding="utf-8")

    module = importlib.import_module("drive_qual.reports.generate")
    output_path = module.generate_report_docx(part_number="69-420", source_root=source_root)
    document = Document(output_path)
    headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style is not None and paragraph.style.name.startswith("Heading")
    ]
    table_headers = [" | ".join(cell.text for cell in table.rows[0].cells) for table in document.tables]

    assert headings[:4] == [
        "Drive Qualification Report.",
        "Drive Info",
        "Qualification Equipment",
        "Test Procedure",
    ]
    assert "Test Results" in headings
    assert "Power Data" in headings
    assert "Compatibility Data" in headings
    assert "Disk Performance" in headings
    assert table_headers[:6] == [
        "Revision | Name | Date | Description",
        "Test | Linux | MacOS | Windows",
        "Test | Linux | MacOS | Windows",
        "Temperature | Read MB/s | Write MB/s",
        "DUT | CDM-R | CDM-W | BM(R) | BM(W) | ATTO-R | ATTO-W",
        "Program | Iterations/Loops | Result",
    ]
    assert "Artifact | Path" not in table_headers
    assert _has_paragraph_between_tables(document, "Windows | ", "Linux | ")
    assert _has_paragraph_between_tables(document, "Linux | ", "MAC | ")
    assert _first_column_is_narrower_than_artifact_column(document.tables[6])


def test_generate_report_docx_embeds_appendix_images_instead_of_paths() -> None:
    from docx import Document

    source_root = Path("tests/.tmp/test_report_generation_images")
    part_dir = source_root / "69-420"
    windows_dir = part_dir / "Windows"
    linux_dir = part_dir / "Linux"
    windows_dir.mkdir(parents=True, exist_ok=True)
    linux_dir.mkdir(parents=True, exist_ok=True)
    report_path = part_dir / "drive_qualification_report_atomic_tests.json"
    report_path.write_text(json.dumps(_report_payload()), encoding="utf-8")
    _write_png(windows_dir / "Padlock DT Inrush Summary.png")
    _write_measurement_csv(windows_dir / "Padlock DT Inrush Summary.csv")
    _write_png(windows_dir / "Padlock DT Max IO Summary.png")
    _write_measurement_csv(windows_dir / "Padlock DT Max IO Summary.csv")
    _write_png(windows_dir / "Padlock DT CrystalDiskMark Performance.png")
    (linux_dir / "Padlock DT Max IO Summary.csv").write_text("time,current\n", encoding="utf-8")

    module = importlib.import_module("drive_qual.reports.generate")
    output_path = module.generate_report_docx(part_number="69-420", source_root=source_root)
    document = Document(output_path)
    table_headers = [" | ".join(cell.text for cell in table.rows[0].cells) for table in document.tables]
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    document_text += "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)

    _assert_appendix_object_layout(document, output_path)
    assert "Artifact | Path" not in table_headers
    assert str(windows_dir) not in document_text
    assert _nested_table_rows(document.tables[6].rows[2].cells[1]) == [
        ["Name", "Measurement", "Accum-Mean", "Accum-Min", "Accum-Max"],
        ["Meas1", "Maximum", "448.48 mA", "444.94 mA", "453.56 mA"],
        ["Meas3", "RMS", "258.60 mA", "258.04 mA", "259.17 mA"],
    ]
    assert _nested_table_columns_evenly_fill_parent(document.tables[6].rows[2].cells[1])
    assert document.tables[7].rows[2].cells[1].text == "Linux\\Padlock DT Max IO Summary.csv"


def _assert_appendix_object_layout(document: Any, output_path: Path) -> None:
    assert len(document.inline_shapes) == EXPECTED_INLINE_IMAGE_COUNT
    assert _embedded_object_count(output_path) == EXPECTED_POWER_OBJECT_COUNT
    assert _media_file_count(output_path) == EXPECTED_MEDIA_FILE_COUNT
    assert _table_cell_object_count(document.tables[6], 1, 0) == 1
    assert _table_cell_object_count(document.tables[6], 2, 0) == 1
    assert _table_cell_object_count(document.tables[6], 1, 1) == 0
    assert _table_cell_object_count(document.tables[6], 2, 1) == 0
    assert _cell_paragraph_texts(document.tables[6], 1, 0)[:2] == ["Inrush Summary", ""]
    assert _cell_paragraph_texts(document.tables[6], 2, 0)[:2] == ["Max IO Summary", ""]
    assert _table_cell_drawing_count(document.tables[6], 2, 1) == 0
    assert _table_cell_drawing_count(document.tables[6], 3, 1) == 1
    assert _embedded_object_payload_names(output_path) == [
        "Padlock DT Inrush Summary.png",
        "Padlock DT Max IO Summary.png",
    ]
    assert _embedded_object_shape_ids(output_path) == ["_x0000_i1009", "_x0000_i1011"]
    assert all(object_id < MAX_WORD_OBJECT_ID for object_id in _embedded_object_ids(output_path))


def test_power_evaluation_applies_max_io_thresholds() -> None:
    result = evaluate_power(
        {
            "Padlock DT": {
                "rms_read_write_current_5v": {
                    "windows": 899,
                    "linux": 900,
                    "macos": 1000,
                }
            }
        }
    )

    statuses = {row.label: row.status for row in result["Padlock DT"]}

    assert statuses["Max I/O RMS current 5V (windows)"] == Status.PASS
    assert statuses["Max I/O RMS current 5V (linux)"] == Status.WARN
    assert statuses["Max I/O RMS current 5V (macos)"] == Status.FAIL


def test_power_evaluation_applies_voltage_and_inrush_rules() -> None:
    result = evaluate_power(
        {
            "Padlock DT": {
                "min_read_write_voltage_5v": {"windows": 4.8, "linux": 4.7, "macos": None},
                "max_inrush_current_5v": {"windows": 900, "linux": 901, "macos": None},
            }
        }
    )

    statuses = {row.label: row.status for row in result["Padlock DT"]}

    assert statuses["Max I/O minimum voltage 5V (windows)"] == Status.PASS
    assert statuses["Max I/O minimum voltage 5V (linux)"] == Status.FAIL
    assert statuses["Max I/O minimum voltage 5V (macos)"] == Status.MISSING
    assert statuses["In-Rush current 5V (windows)"] == Status.PASS
    assert statuses["In-Rush current 5V (linux)"] == Status.WARN


def test_temperature_evaluation_marks_zero_and_errors_red() -> None:
    result = evaluate_temperature(
        {
            "Padlock DT": {
                "performance": {
                    "20c": {"read_mb_s": 100, "write_mb_s": 0},
                    "30c": {"read_mb_s": None, "write_mb_s": None, "error": "benchmark failed"},
                }
            }
        }
    )

    statuses = {row.label: row.status for row in result["Padlock DT"]}

    assert statuses["20c read mb s"] == Status.PASS
    assert statuses["20c write mb s"] == Status.FAIL
    assert statuses["30c read mb s"] == Status.FAIL
    assert statuses["30c write mb s"] == Status.FAIL


def test_temperature_material_coverage_marks_untested_like_products_not_applicable() -> None:
    result = evaluate_temperature(
        {
            "Padlock DT": {"performance": {"20c": {"read_mb_s": 100, "write_mb_s": 100}}},
            "ASK3": {"performance": {"20c": {"read_mb_s": None, "write_mb_s": None}}},
        }
    )

    assert result["ASK3"][0].status == Status.NOT_APPLICABLE
    assert "Padlock DT" in result["ASK3"][0].reason


def test_case_material_classification_defaults_to_plastic() -> None:
    assert case_material_for_product("Fortress L3") == "aluminum"
    assert case_material_for_product("Padlock DT FIPS") == "aluminum"
    assert case_material_for_product("Padlock NVX") == "plastic"


def _write_png(path: Path) -> None:
    from PIL import Image

    seed = sum(path.name.encode("utf-8"))
    image = Image.new("RGB", (320, 320))
    pixels = [
        ((x * 13 + y * 7 + seed) % 256, (x * 5 + y * 17 + seed) % 256, (x * y + seed) % 256)
        for y in range(320)
        for x in range(320)
    ]
    image.putdata(cast(Any, pixels))
    image.save(path)


def _write_measurement_csv(path: Path) -> None:
    path.write_text(
        "\n".join(
            [
                "TekScope,Version 2.0.3",
                "",
                "Measurement Results",
                "Name,Measurement,Label,Source,Mean',Accum-Mean,Accum-Min,Accum-Max,Accum-Pk-Pk,"
                "Accum-Std Dev,Accum-Population",
                'Meas1,Maximum,Maximum," Ch 4 ",448.62 mA,448.48 mA,444.94 mA,453.56 mA,8.6250 mA,1.5484 mA,132',
                'Meas3,RMS,RMS," Ch 4 ",258.70 mA,258.60 mA,258.04 mA,259.17 mA,1.1256 mA,212.51 uA,132',
                'Meas9,Peak,Peak," Ch 4 ",999.00 mA,999.00 mA,998.00 mA,1000.00 mA,2.0000 mA,1.0000 mA,132',
            ]
        ),
        encoding="utf-8",
    )


def _table_cell_drawing_count(document_table: Any, row_index: int, cell_index: int) -> int:
    return len(_cell_xml(document_table, row_index, cell_index).xpath(".//w:drawing"))


def _table_cell_object_count(document_table: Any, row_index: int, cell_index: int) -> int:
    return len(_cell_xml(document_table, row_index, cell_index).xpath(".//w:object"))


def _cell_xml(document_table: Any, row_index: int, cell_index: int) -> Any:
    return document_table.rows[row_index].cells[cell_index]._tc


def _cell_paragraph_texts(document_table: Any, row_index: int, cell_index: int) -> list[str]:
    return [paragraph.text for paragraph in document_table.rows[row_index].cells[cell_index].paragraphs]


def _embedded_object_count(path: Path) -> int:
    with zipfile.ZipFile(path) as docx_zip:
        return len([name for name in docx_zip.namelist() if name.startswith("word/embeddings/oleObject")])


def _media_file_count(path: Path) -> int:
    with zipfile.ZipFile(path) as docx_zip:
        return len([name for name in docx_zip.namelist() if name.startswith("word/media/")])


def _embedded_object_payload_names(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as docx_zip:
        object_names = [name for name in docx_zip.namelist() if name.startswith("word/embeddings/oleObject")]
        return [_ole10_native_label(docx_zip.read(name)) for name in object_names]


def _embedded_object_shape_ids(path: Path) -> list[str]:
    return [match.split('"')[1] for match in _embedded_object_xml_matches(path, 'ShapeID="')]


def _embedded_object_ids(path: Path) -> list[int]:
    return [int(match.split('"')[1].removeprefix("_")) for match in _embedded_object_xml_matches(path, 'ObjectID="')]


def _embedded_object_xml_matches(path: Path, marker: str) -> list[str]:
    with zipfile.ZipFile(path) as docx_zip:
        xml = docx_zip.read("word/document.xml").decode("utf-8")
    return [token for token in xml.split() if token.startswith(marker)]


def _ole10_native_label(blob: bytes) -> str:
    marker = "\x01Ole10Native".encode("utf-16le")
    directory_offset = blob.find(marker)
    assert directory_offset >= 0
    stream_start = int.from_bytes(blob[directory_offset + 116 : directory_offset + 120], "little", signed=True)
    sector_offset = (stream_start + 1) * 512
    stream = blob[sector_offset : sector_offset + 512]
    label_start = 6
    label_end = stream.index(b"\x00", label_start)
    return stream[label_start:label_end].decode("utf-8")


def _nested_table_rows(cell: Any) -> list[list[str]]:
    table = cell.tables[0]
    return [[nested_cell.text for nested_cell in row.cells] for row in table.rows]


def _nested_table_columns_evenly_fill_parent(cell: Any) -> bool:
    widths = [nested_cell.width for nested_cell in cell.tables[0].rows[0].cells]
    return all(width is not None for width in widths) and len(set(widths)) == 1 and sum(widths) < cell.width


def _first_column_is_narrower_than_artifact_column(table: Any) -> bool:
    first_width = table.rows[0].cells[0].width
    second_width = table.rows[0].cells[1].width
    return first_width is not None and second_width is not None and first_width * 2 <= second_width


def _has_paragraph_between_tables(document: Any, first_header: str, second_header: str) -> bool:
    labels = _body_block_labels(document)
    first_index = labels.index(f"table:{first_header}")
    second_index = labels.index(f"table:{second_header}")
    return "paragraph" in labels[first_index + 1 : second_index]


def _body_block_labels(document: Any) -> list[str]:
    table_headers = iter(" | ".join(cell.text for cell in table.rows[0].cells) for table in document.tables)
    labels: list[str] = []
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}tbl"):
            labels.append(f"table:{next(table_headers)}")
        elif child.tag.endswith("}p"):
            labels.append("paragraph")
    return labels


def _report_payload() -> dict[str, Any]:
    return {
        "drive_info": {
            "apricorn_part_number": "69-420",
            "manufacturer": "Apricorn",
            "manufacturer_part_number": "ABC123",
            "serial_number": "SER123",
            "capacity": "1TB",
            "firmware": "1.0",
            "form_factor": "2.5",
            "interface": "USB",
            "technology": "SSD",
        },
        "equipment": {
            "scope": {"model": "Tektronix MSO54", "version": "2.0.3", "serial_number": "B013976"},
            "probe_current": {"model": "TCP202A", "channel": "4", "serial_number": "C004510"},
            "probe_voltage": {"model": "TPP0500B", "channel": "2", "serial_number": "C166742"},
            "windows_host": {
                "hardware": "ASUS PRIME Z270-K, i5-7400K",
                "os_version": "Windows 10",
                "software": [{"name": "CrystalDiskMark", "version": "7.0.0"}, {"name": "ATTO", "version": "4.0"}],
            },
            "linux_host": {"hardware": "NUC", "os_version": "Ubuntu", "software": [{"name": "Disks", "version": None}]},
            "macos_host": {
                "hardware": "Mac Mini",
                "os_version": "OS 15",
                "software": [{"name": "Blackmagic Disk Speed Test", "version": "4.2"}],
            },
            "dut": {"Padlock DT": {"serial_number": "DUT123"}},
        },
        "compatibility": {
            "recognized_by_os": {"linux": True, "macos": True, "windows": True},
            "hot_pluggable": {"linux": True, "macos": True, "windows": True},
        },
        "power": {
            "Padlock DT": {
                "max_inrush_current_5v": {"linux": 700, "macos": 800, "windows": 750},
                "max_read_write_current_5v": {"linux": 850, "macos": 860, "windows": 870},
                "rms_read_write_current_5v": {"linux": 500, "macos": 510, "windows": 520},
            }
        },
        "performance": {
            "Padlock DT": {
                "Windows": {
                    "CrystalDiskMark": {"read": 350.93, "write": 345.04},
                    "ATTO": {"read": 334.95, "write": 318.74},
                },
                "macOS": {"Blackmagic Disk Speed Test": {"read": 293.2, "write": 844.2}},
            }
        },
        "temperature": {"Padlock DT": {"performance": {"-35c": {"read_mb_s": 107.59, "write_mb_s": 109.31}}}},
        "compliance": {
            "usb_if_msc_iterations": 3,
            "usb_if_msc_result": "Pass",
            "disk_tester_reliability_iterations": 3,
            "disk_tester_reliability_result": "Pass",
        },
    }
