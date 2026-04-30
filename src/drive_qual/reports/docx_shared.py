from __future__ import annotations

import re
from collections.abc import Callable, Iterable
from pathlib import Path
from typing import Any

from drive_qual.reports.constants import APPENDIX_MEASUREMENT_COLUMN_WIDTH_FRACTION, EMU_PER_TWIP, STATUS_COLORS
from drive_qual.reports.evaluation import Status


def _load_docx_tools() -> tuple[Any, Any, Callable[[Any, Status], None]]:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError(
            "Word report generation requires python-docx. Install project dependencies with `uv sync`."
        ) from exc

    def shade_cell(cell: Any, status: Status) -> None:
        _shade_cell(cell, status)

    return Document, Inches, shade_cell


def _shade_cell(cell: Any, status: Status) -> None:
    OxmlElement, qn = _load_docx_xml_tools()
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), STATUS_COLORS[status])
    tc_pr.append(shading)


def _set_margins(document: Any, inches: Any) -> None:
    for section in document.sections:
        section.left_margin = inches(0.6)
        section.right_margin = inches(0.6)
        section.top_margin = inches(0.5)
        section.bottom_margin = inches(0.5)


def _add_header_logo(document: Any, inches: Any) -> None:
    logo_path = Path(__file__).parent / "logo.png"
    if not logo_path.exists():
        return
    section = document.sections[0]
    section.different_first_page_header_footer = True
    section.header_distance = inches(0.5)
    header = section.first_page_header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(str(logo_path), height=inches(0.5))


def _add_title(document: Any) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    heading = _add_heading(document, "Drive Qualification Report", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.runs[0]
    run.font.name = "Times New Roman"
    run.font.size = Pt(24)
    run.underline = True
    document.add_paragraph()


def _add_footer(document: Any) -> None:
    for section in document.sections:
        section.different_first_page_header_footer = True
        _populate_footer(section.first_page_footer, section)
        _populate_footer(section.footer, section)


def _populate_footer(footer: Any, section: Any) -> None:
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Pt, RGBColor

    paragraph = footer.paragraphs[0]
    paragraph.clear()
    available_width = section.page_width - section.left_margin - section.right_margin
    paragraph.paragraph_format.tab_stops.add_tab_stop(int(available_width / 2), WD_TAB_ALIGNMENT.CENTER)
    paragraph.paragraph_format.tab_stops.add_tab_stop(int(available_width), WD_TAB_ALIGNMENT.RIGHT)

    left_run = paragraph.add_run("Drive Qualification Report.docx")
    confidential_run = paragraph.add_run("\tApricorn Confidential")
    confidential_run.font.bold = True
    confidential_run.font.color.rgb = RGBColor(255, 0, 0)
    runs = [
        left_run,
        confidential_run,
        paragraph.add_run("\tPage "),
    ]
    _add_field(paragraph, "PAGE")
    runs.append(paragraph.add_run(" of "))
    _add_field(paragraph, "NUMPAGES")
    for run in runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)


def _add_field(paragraph: Any, instruction: str) -> None:
    OxmlElement, qn = _load_docx_xml_tools()
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    for tag, attrs in (
        ("w:fldChar", {"w:fldCharType": "begin"}),
        ("w:instrText", {"xml:space": "preserve", "text": f" {instruction} "}),
        ("w:fldChar", {"w:fldCharType": "separate"}),
        ("w:t", {"text": "1"}),
        ("w:fldChar", {"w:fldCharType": "end"}),
    ):
        element = OxmlElement(tag)
        for key, value in attrs.items():
            if key == "text":
                element.text = value
            elif key == "xml:space":
                element.set("{http://www.w3.org/XML/1998/namespace}space", value)
            else:
                element.set(qn(key), value)
        run._r.append(element)


def _table(document: Any, headers: Iterable[str]) -> Any:
    header_list = list(headers)
    table = document.add_table(rows=1, cols=len(header_list))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, header_list, strict=True):
        cell.text = header
    return table


def _set_table_column_widths(table: Any, widths: list[Any]) -> None:
    OxmlElement, qn = _load_docx_xml_tools()
    table.autofit = False
    _set_table_preferred_width(table, sum(widths), OxmlElement, qn)
    _set_table_grid(table, widths, OxmlElement, qn)
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = width
            _set_cell_preferred_width(row.cells[index], width, OxmlElement, qn)


def _load_docx_xml_tools() -> tuple[Any, Any]:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    return OxmlElement, qn


def _set_table_preferred_width(table: Any, width: Any, oxml_element: Any, qn: Any) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = oxml_element("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(_twips(width)))


def _set_table_grid(table: Any, widths: list[Any], oxml_element: Any, qn: Any) -> None:
    tbl = table._tbl
    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    tbl_grid = oxml_element("w:tblGrid")
    for width in widths:
        grid_col = oxml_element("w:gridCol")
        grid_col.set(qn("w:w"), str(_twips(width)))
        tbl_grid.append(grid_col)
    tbl.insert(1, tbl_grid)


def _set_cell_preferred_width(cell: Any, width: Any, oxml_element: Any, qn: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = oxml_element("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:type"), "dxa")
    tc_width.set(qn("w:w"), str(_twips(width)))


def _twips(width: Any) -> int:
    return max(int(int(width) / EMU_PER_TWIP), 1)


def _center_cell_paragraphs(cell: Any) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _center_table_column(table: Any, column_index: int) -> None:
    for row in table.rows:
        _center_cell_paragraphs(row.cells[column_index])


def _center_table_columns(table: Any, column_indexes: Iterable[int]) -> None:
    for column_index in column_indexes:
        _center_table_column(table, column_index)


def _minimize_empty_cell_paragraphs(cell: Any) -> None:
    for paragraph in cell.paragraphs:
        if paragraph.text or paragraph._p.xpath(".//w:drawing | .//w:object"):
            continue
        _minimize_empty_paragraph(paragraph)


def _minimize_trailing_empty_cell_paragraph(cell: Any) -> None:
    if not cell.paragraphs:
        return
    paragraph = cell.paragraphs[-1]
    if paragraph.text or paragraph._p.xpath(".//w:drawing | .//w:object"):
        return
    _minimize_empty_paragraph(paragraph)


def _minimize_empty_paragraph(paragraph: Any) -> None:
    from docx.shared import Pt

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(1)
    if not paragraph.runs:
        paragraph.add_run()
    for run in paragraph.runs:
        run.font.size = Pt(1)


def _add_minimized_empty_paragraph(cell: Any) -> None:
    paragraph = cell.add_paragraph("")
    _minimize_empty_paragraph(paragraph)


def _add_nested_table(cell: Any, rows: list[list[str]], table_width: Any) -> Any:
    normalized_rows = _normalize_table_rows(rows)
    table = cell.add_table(rows=1, cols=len(normalized_rows[0]))
    table.style = "Table Grid"
    for index, header in enumerate(normalized_rows[0]):
        table.rows[0].cells[index].text = header
    for values in normalized_rows[1:]:
        row = table.add_row().cells
        for index, value in enumerate(values):
            row[index].text = value
    column_width = int((table_width // len(normalized_rows[0])) * APPENDIX_MEASUREMENT_COLUMN_WIDTH_FRACTION)
    _set_table_column_widths(table, [column_width] * len(normalized_rows[0]))
    return table


def _remove_extra_empty_paragraphs(cell: Any) -> None:
    non_empty_paragraphs = [paragraph for paragraph in cell.paragraphs if paragraph.text]
    if not non_empty_paragraphs:
        return
    for paragraph in list(cell.paragraphs):
        if paragraph.text:
            continue
        paragraph._element.getparent().remove(paragraph._element)


def _normalize_table_rows(rows: list[list[str]]) -> list[list[str]]:
    width = max(len(row) for row in rows)
    return [row + ([""] * (width - len(row))) for row in rows]


def _add_picture_to_paragraph(paragraph: Any, artifact: Path, *, width: Any) -> None:
    try:
        paragraph.add_run().add_picture(str(artifact), width=width)
    except Exception:
        paragraph.add_run(str(artifact.name))


def _add_heading(document: Any, text: str, level: int) -> Any:
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Times New Roman"
    return heading


def _list_line(document: Any, text: str, level: int = 0) -> None:
    from docx.shared import Inches

    p = document.add_paragraph(text, style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))


def _field(data: object, key: str) -> str:
    if not isinstance(data, dict):
        return ""
    return _format_value(data.get(key))


def _format_value(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        return f"{value:.2f}".rstrip("0").rstrip(".")
    if isinstance(value, list):
        return ", ".join(_format_value(item) for item in value)
    if isinstance(value, dict):
        return "; ".join(f"{_label(str(key))}: {_format_value(val)}" for key, val in value.items())
    return str(value)


def _label(value: str) -> str:
    return value.replace("_", " ").title()


def _join_present(*values: object) -> str:
    return ", ".join(str(value) for value in values if value)


def _normalize_text(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", value.casefold()).strip()


def _to_float(value: object) -> float | None:
    if value is None:
        return None
    try:
        return float(str(value).strip())
    except ValueError:
        return None
